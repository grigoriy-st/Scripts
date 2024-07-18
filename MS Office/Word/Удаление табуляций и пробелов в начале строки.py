import docx

doc = docx.Document('Комбинация клавиш.docx')
document = docx.Document()


for paragraph in doc.paragraphs:
    if paragraph.text.startswith('\t') or paragraph.text.startswith(' '):
        document.add_paragraph(paragraph.text.lstrip("\t").lstrip(" "))
    else:
        document.add_paragraph(paragraph.text)
        
document.save("New.docx")
